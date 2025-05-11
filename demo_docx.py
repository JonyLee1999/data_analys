from docx import Document
import docx.opc.constants
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

# 创建新文档
document = Document()

# ========== 自定义样式 ==========
# 创建标题样式
styles = document.styles
title_style = styles.add_style('MyTitle', WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = '微软雅黑'
title_style.font.size = Pt(24)
title_style.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 创建自定义列表样式
list_style = styles.add_style('MyList', WD_STYLE_TYPE.PARAGRAPH)
list_style.paragraph_format.left_indent = Inches(0.5)
list_style.font.name = 'Consolas'

# ========== 文档属性 ==========
document.core_properties.author = "AI助手"
document.core_properties.comments = "使用python-docx创建的复杂示例文档"
document.core_properties.keywords = "Python, DOCX, 示例"

# ========== 页眉页脚 ==========
section = document.sections[0]
header = section.header
footer = section.footer

# 页眉内容
header_para = header.paragraphs[0]
header_para.text = "机密文档 - 禁止未经授权的传播\t\t生成日期：" + datetime.now().strftime("%Y-%m-%d")
header_para.style = document.styles['Header']

# 页脚页码
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    page_run = paragraph.add_run()
    page_run.text = "第 "
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)
    
    page_run = paragraph.add_run()
    page_run.text = " 页，共 "
    
    total_run = paragraph.add_run()
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    total_run._r.append(fldChar3)
    total_run._r.append(instrText2)
    total_run._r.append(fldChar4)
    
    page_run = paragraph.add_run()
    page_run.text = " 页"

# 添加页码
footer_para = footer.paragraphs[0]
add_page_number(footer_para)

# ========== 封面页 ==========
# 封面标题
cover = document.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_after = Pt(72)
run = cover.add_run("项目报告\n")
run.font.size = Pt(36)
run.bold = True
run = cover.add_run("Python-docx高级示例")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 封面信息表格
cover_table = document.add_table(rows=3, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_table.style = 'LightShading-Accent1'
cells = cover_table.rows[0].cells
cells[0].text = "项目名称"
cells[1].text = "Python-docx功能演示"

cells = cover_table.rows[1].cells
cells[0].text = "作者"
cells[1].text = "AI助手"

cells = cover_table.rows[2].cells
cells[0].text = "版本"
cells[1].text = "v2.0"

document.add_page_break()

# ========== 正文内容 ==========
# 使用自定义标题样式
document.add_paragraph('1. 功能演示', style='MyTitle')

# 复杂段落格式
p = document.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
p.add_run("这是首行缩进的段落。").bold = True
p.add_run(" 这里演示不同字体：")
run = p.add_run("华文行楷")
run.font.name = '华文行楷'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文行楷')

# 添加分栏（通过分节符）
document.add_section()
document.sections[-1].start_type  # 新分节符会自动添加

# 多级列表
document.add_paragraph("一级列表", style='List Number')
document.add_paragraph("二级列表", style='List Number 2')
document.add_paragraph("三级列表", style='List Number 3')
document.add_paragraph("另一个一级列表", style='List Number')

# 带格式的代码块
code_style = styles.add_style('CodeText', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(11)
code_style.paragraph_format.left_indent = Inches(0.5)
code_style.paragraph_format.space_before = Pt(6)
code_style.paragraph_format.space_after = Pt(6)
code_text = '''def hello_world():
    print("Hello World!")
'''
document.add_paragraph(code_text, style='CodeText')

# 复杂表格
table = document.add_table(rows=3, cols=3)
table.style = 'MediumGrid3-Accent1'
table.autofit = False

# 设置列宽
for row in table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(3)

# 合并单元格
table.cell(0, 0).merge(table.cell(0, 2))
table.cell(0, 0).text = "合并的标题单元格"
table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表格内容
data = [
    ["序号", "项目", "描述"],
    [1, "Python", "高级编程语言"],
    [2, "DOCX", "Word文档处理"]
]

for row_idx in range(3):
    for col_idx in range(3):
        cell = table.cell(row_idx, col_idx)
        cell.text = str(data[row_idx][col_idx])
        if row_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

# ========== 图文混排 ==========
document.add_heading('2. 图片与文字环绕', level=1)
p = document.add_paragraph()
run = p.add_run()
run.add_picture('sales_report.png', width=Inches(2))  # 使用已有的图片文件
run.add_text('  这是文字环绕图片的示例。图片被设置为浮动格式，文字可以环绕在图片周围。此处演示如何在段落中添加图片并进行文字环绕。')

# ========== 高级格式 ==========
document.add_heading('3. 高级格式设置', level=1)

# 带边框段落
bordered_para = document.add_paragraph()
bordered_para.paragraph_format.border_bottom = True
bordered_para.paragraph_format.border_top = True
bordered_para.add_run("带上下边框的段落")

# 文字方向
vertical_para = document.add_paragraph()
vertical_para.paragraph_format.right_indent = Inches(2)
vertical_text = vertical_para.add_run("竖排文字")
vertical_text.font.name = '@微软雅黑'

# 设置文字方向
rPr = vertical_text._element.get_or_add_rPr()
textDirection = OxmlElement('w:textDirection')
textDirection.set(qn('w:val'), 'tbRl')  # top to bottom, right to left
rPr.append(textDirection)

# 超链接
from docx.shared import RGBColor
paragraph = document.add_paragraph()
paragraph.add_run("访问Python官网：")

# 创建超链接
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), 'rId1000')  # 任意唯一ID
tooltip = OxmlElement('w:tooltip')
tooltip.set(qn('w:val'), "点击访问")
hyperlink.append(tooltip)

# 创建链接文本运行对象
new_run = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')

# 设置链接样式
color = OxmlElement('w:color')
color.set(qn('w:val'), '0000FF')  # 蓝色
rPr.append(color)

# 添加下划线
u = OxmlElement('w:u')
u.set(qn('w:val'), 'single')
rPr.append(u)

new_run.append(rPr)
t = OxmlElement('w:t')
t.text = 'Python'
new_run.append(t)

hyperlink.append(new_run)
paragraph._p.append(hyperlink)

# 添加关系
rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
target_url = "https://www.python.org"
rid = paragraph.part.relate_to(target_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
hyperlink.set(qn('r:id'), rid)

# ========== 保存文档 ==========
document.add_page_break()
document.save('advanced_demo.docx')